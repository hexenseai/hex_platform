from django.db import models
from django.contrib.auth.models import User
from .utils import avatar_upload_path, company_logo_upload_path
import uuid
from qdrant_client import QdrantClient
from sentence_transformers import SentenceTransformer
import os
from qdrant_client.http import models as qdrant_models
import tiktoken
import mimetypes
import pdfplumber
import docx
import open_clip
from PIL import Image
import io
import pandas as pd

# CLIP model yüklemesi (ilk kullanımda yüklenir)
_clip_model = None
_clip_preprocess = None
_clip_tokenizer = None

def get_clip_model():
    global _clip_model, _clip_preprocess, _clip_tokenizer
    if _clip_model is None:
        _clip_model, _, _clip_preprocess = open_clip.create_model_and_transforms('ViT-B-32', pretrained='openai')
        _clip_tokenizer = open_clip.get_tokenizer('ViT-B-32')
    return _clip_model, _clip_preprocess

class Company(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    name = models.CharField(max_length=255)
    description = models.TextField(blank=True)
    address = models.TextField(blank=True)
    logo = models.ImageField(upload_to=company_logo_upload_path, blank=True, null=True)
    openai_api_key = models.CharField(max_length=255, blank=True, null=True)
    gemini_api_key = models.CharField(max_length=255, blank=True, null=True)
    claude_api_key = models.CharField(max_length=255, blank=True, null=True)
    deepseek_api_key = models.CharField(max_length=255, blank=True, null=True)
    azure_api_key = models.CharField(max_length=255, blank=True, null=True)
    azure_endpoint = models.CharField(max_length=255, blank=True, null=True)
    
    def __str__(self):
        return self.name
    
    class Meta:
        verbose_name = "Company"
        verbose_name_plural = "Organisation: Companies"

    
class Department(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    company = models.ForeignKey(Company, on_delete=models.CASCADE, related_name='departments')
    name = models.CharField(max_length=150)
    description = models.TextField(blank=True)
    parent = models.ForeignKey('self', null=True, blank=True, on_delete=models.SET_NULL, related_name='sub_departments')

    def __str__(self):
        return f"{self.name} ({self.company.name})"
    
    class Meta:
        verbose_name = "Department"
        verbose_name_plural = "Organisation: Departments"
    

class Role(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    department = models.ForeignKey(Department, on_delete=models.CASCADE, related_name='roles')
    name = models.CharField(max_length=150)
    description = models.TextField(blank=True)

    def __str__(self):
        return f"{self.name} ({self.department.name if self.department else 'No Department'})"
    
    class Meta:
        verbose_name = "Role"
        verbose_name_plural = "Organisation: Roles"


class UserProfile(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='profiles')
    company = models.ForeignKey(Company, null=True, blank=True, on_delete=models.SET_NULL, related_name='users')
    department = models.ForeignKey(Department, null=True, blank=True, on_delete=models.SET_NULL, related_name='users')
    role = models.ForeignKey(Role, null=True, blank=True, on_delete=models.SET_NULL, related_name='users')
    phone_number = models.CharField(max_length=20, blank=True)
    avatar = models.ImageField(upload_to=avatar_upload_path, blank=True, null=True)
    gpt_preferences = models.TextField(blank=True, help_text="GPT kullanım tercihlerini açıklayın.")
    work_experience_notes = models.TextField(blank=True, help_text="İş geçmişiniz veya uzmanlık alanlarınız hakkında bilgi girin.")
    is_current = models.BooleanField(default=False, help_text="Bu profil kullanıcının aktif profili mi?")
    company_admin = models.BooleanField(default=False, help_text="Firma bilgilerini ve kullanıcıları yönetme yetkisi")

    def __str__(self):
        return f"{self.user.username} ({self.company.name if self.company else ''})"
    
    class Meta:
        verbose_name = "User Profile"
        verbose_name_plural = "Organisation: User Profiles"


class Conversation(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    user_profile = models.ForeignKey('UserProfile', null=True, blank=True, on_delete=models.SET_NULL, related_name='conversations')
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    context = models.TextField(blank=True, null=True)
    is_active = models.BooleanField(default=True)
    
    # İleride GPT paket bilgisi eklenebilir
    def __str__(self):
        return self.title or f"Conversation {self.pk}"

    class Meta:
        verbose_name = "Conversation"
        verbose_name_plural = "Agent Management: Conversations"

    def save(self, *args, **kwargs):
        is_new = not self.pk  # Check if this is a new record
        if not is_new:
            from hexense_core.semantic import delete_from_qdrant
            delete_from_qdrant("conversations", [str(self.id)])
        super().save(*args, **kwargs)
        from hexense_core.semantic import get_embedding, add_to_qdrant
        embedding = get_embedding(self.context or "")
        payload = {
            "conversation_id": str(self.id),
            "is_active": self.is_active,
            "created_at": self.created_at.isoformat(),
        }
        add_to_qdrant("conversations", self.context or "", payload, vector=embedding, point_id=str(self.id))

    def delete(self, *args, **kwargs):
        from hexense_core.semantic import delete_from_qdrant
        delete_from_qdrant("conversations", [str(self.id)])
        super().delete(*args, **kwargs)


class Message(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    conversation = models.ForeignKey(Conversation, on_delete=models.CASCADE, related_name='messages')
    SENDER_CHOICES = [
        ('user', 'User'),
        ('assistant', 'Assistant'),
        ('tool', 'Tool'),
    ]
    sender = models.CharField(max_length=10, choices=SENDER_CHOICES)
    content = models.TextField()
    gpt_package = models.ForeignKey('GptPackage', null=True, blank=True, on_delete=models.SET_NULL)
    actions = models.JSONField(blank=True, null=True)
    timestamp = models.DateTimeField(auto_now_add=True)
    is_active = models.BooleanField(default=True)

    def __str__(self):
        return f"{self.sender}: {self.content[:50]}"
    
    class Meta:
        verbose_name = "Message"
        verbose_name_plural = "Agent Management: Messages"

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
        from hexense_core.semantic import get_embedding, add_to_qdrant
        embedding = get_embedding(self.content)
        payload = {
            "conversation_id": str(self.conversation_id),
            "gpt_package_id": str(self.gpt_package_id) if self.gpt_package_id else None,
            "timestamp": self.timestamp.isoformat(),
            "message_id": str(self.id),
            "sender": self.sender,
            "is_active": self.is_active,
        }
        add_to_qdrant("messages", self.content, payload, vector=embedding, point_id=str(self.id))

    def delete(self, *args, **kwargs):
        from hexense_core.semantic import delete_from_qdrant
        delete_from_qdrant("messages", [str(self.id)])
        super().delete(*args, **kwargs)

    def set_active(self, is_active: bool):
        self.is_active = is_active
        self.save()
        from hexense_core.semantic import update_qdrant_metadata
        update_qdrant_metadata("messages", str(self.id), {"is_active": is_active})


import uuid
from django.db import models


class GptModel(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    key = models.CharField(max_length=50, unique=True)

    provider = models.CharField(
        max_length=20,
        choices=[
            ('openai', 'OpenAI'),
            ('anthropic', 'Anthropic Claude'),
            ('gemini', 'Google Gemini'),
        ]
    )
    name = models.CharField(max_length=100)  # Örnek: gpt-4o, claude-3-opus, gemini-pro
    model_family = models.CharField(max_length=50, blank=True, help_text="GPT-4, Claude 3, Gemini 2.5 gibi")
    description = models.TextField(blank=True)

    is_active = models.BooleanField(default=True)
    is_local = models.BooleanField(default=False, help_text="Model yerel olarak mı çalıştırılıyor?")
    local_path = models.CharField(max_length=255, blank=True, null=True, help_text="Hugging Face veya dosya yolu")

    # Teknik özellikler
    context_window = models.IntegerField(null=True, blank=True, help_text="Maksimum bağlam penceresi (token)")
    supports_function_call = models.BooleanField(default=False)
    supports_multimodal = models.BooleanField(default=False)
    supports_tool_use = models.BooleanField(default=False)

    # Giriş / çıkış yetenekleri
    input_types = models.JSONField(
        default=list,
        blank=True,
        help_text="Desteklenen input türleri: ['text', 'image', 'file_upload', 'url', 'audio', 'video', 'pdf', 'table']"
    )
    output_types = models.JSONField(
        default=list,
        blank=True,
        help_text="Beklenen çıktı türleri: ['text', 'image', 'json', 'table', 'audio']"
    )

    # Kullanım senaryoları ve kabiliyetler
    capabilities = models.JSONField(default=list, blank=True, help_text="['text', 'image', 'audio', 'code']")
    optimized_for = models.JSONField(default=list, blank=True, help_text="['general', 'code', 'reasoning', 'speed', 'cost']")

    # Fiyatlandırma
    pricing_per_1k = models.JSONField(
        default=dict,
        blank=True,
        help_text='{"input": 0.01, "output": 0.03} (USD per 1K tokens)'
    )

    release_date = models.DateField(null=True, blank=True)

    def get_token_cost(self, input_tokens: int, output_tokens: int) -> float:
        in_rate = self.pricing_per_1k.get("input", 0)
        out_rate = self.pricing_per_1k.get("output", 0)
        return round((input_tokens / 1000 * in_rate) + (output_tokens / 1000 * out_rate), 4)

    def __str__(self):
        return f"{self.name} ({self.provider})"

    class Meta:
        verbose_name = "Model"
        verbose_name_plural = "Agent Management: Models"



class GptPackageGroup(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    key = models.CharField(max_length=50, unique=True)
    name = models.CharField(max_length=100)
    description = models.TextField(blank=True)

    def __str__(self):
        return self.name
    
    class Meta:
        verbose_name = "Package"
        verbose_name_plural = "Agent Management: Packages"


class GptPackage(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    group = models.ForeignKey('GptPackageGroup', on_delete=models.CASCADE, related_name='packages')
    key = models.CharField(max_length=50, unique=True)
    name = models.CharField(max_length=100)
    description = models.TextField(blank=True)
    model = models.ForeignKey('GptModel', null=True, blank=True, on_delete=models.PROTECT, related_name='packages')
    system_prompt = models.TextField()
    services = models.ManyToManyField('GptService', blank=True, related_name='packages')
    allowed_roles = models.ManyToManyField('Role', blank=True, related_name='gpt_packages')
    is_default = models.BooleanField(default=False)
    
    include_company_info = models.BooleanField(default=False, help_text="GPT'ye firma bilgilerini ekle")
    include_personal_info = models.BooleanField(default=False, help_text="GPT'ye kişisel bilgileri ekle")

    def __str__(self):
        return f"{self.name} ({self.group.name})"
    
    class Meta:
        verbose_name = "Agent"
        verbose_name_plural = "Agent Management: Agents"

    def build_system_prompt(self, user_profile=None):
        prompt_parts = []
        # Firma bilgisi ekle
        if self.include_company_info and self.model and self.model.packages.exists():
            # GptModel -> allowed_roles -> department -> company
            # Veya doğrudan bir company örneği ile çağrılabilir
            company = None
            # allowed_roles üzerinden bir company bulmaya çalış
            role = self.allowed_roles.first()
            if role and role.department and role.department.company:
                company = role.department.company
            if company:
                prompt_parts.append(f"Şirket Bilgisi: {company.name}\nAçıklama: {company.description}")
        # Kişisel bilgi ekle
        if self.include_personal_info and user_profile:
            prompt_parts.append(f"Kullanıcı: {user_profile.user.username}\nUzmanlık: {user_profile.work_experience_notes}")
        # Ana prompt
        prompt_parts.append(self.description)
        return "\n\n".join([p for p in prompt_parts if p])

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
        # Qdrant koleksiyonu oluşturulmamışsa oluştur
        # Koleksiyon kontrolü semantic.py'ye taşındı
        ensure_collection_exists("gpt_packages", vector_size=384)
        # Embedding oluştur
        description = f"{self.name}: {self.description}"
        payload = {
            "id": str(self.id),
            "gpt_package_id": str(self.id),
            "group_id": str(self.group.id),
            "group_name": self.group.name,
            "group_key": self.group.key,
            "name": self.name,
            "key": self.key
        }
        add_to_qdrant("gpt_packages", description, payload)

    def delete(self, *args, **kwargs):
        # Qdrant'tan sil
        delete_from_qdrant("gpt_packages", [str(self.id)])
        super().delete(*args, **kwargs)


class GptService(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    group = models.ForeignKey('GptPackageGroup', on_delete=models.CASCADE, related_name='services')

    key = models.CharField(max_length=100, unique=True)
    name = models.CharField(max_length=150)
    description = models.TextField()

    input_schema = models.JSONField(help_text="Beklenen parametreler ve açıklamaları", blank=True, default=dict)
    output_schema = models.JSONField(blank=True, null=True, help_text="Yanıt formatı (isteğe bağlı)")

    function_name = models.CharField(
        max_length=100,
        default="call_service",
        help_text="tools.py içinde çağrılacak fonksiyonun adı"
    )
    default_params = models.JSONField(
        blank=True,
        null=True,
        help_text="Fonksiyona eklenecek sabit parametreler (isteğe bağlı)"
    )

    is_active = models.BooleanField(default=True)

    def __str__(self):
        return f"{self.name} ({self.group.name})"
    
    class Meta:
        verbose_name = "Service"
        verbose_name_plural = "Agent Management: Services"


class GptPackageFile(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    gpt_package = models.ForeignKey('GptPackage', on_delete=models.CASCADE, related_name='files')
    file = models.FileField(upload_to='gpt_package_files/')
    description = models.CharField(max_length=255, blank=True)
    uploaded_at = models.DateTimeField(auto_now_add=True)
    uploaded_by = models.ForeignKey(User, null=True, blank=True, on_delete=models.SET_NULL, related_name='uploaded_gpt_package_files')

    def __str__(self):
        return f"{self.file.name} ({self.gpt_package.name})"

    class Meta:
        verbose_name = "Gpt Package File"
        verbose_name_plural = "Agent Management: Gpt Package Files"

    def get_file_content(self):
        ext = os.path.splitext(self.file.name)[1].lower()
        self.file.seek(0)
        if ext == '.pdf':
            return self._extract_pdf_text(), self._extract_pdf_images()
        elif ext == '.docx':
            return self._extract_docx_text(), self._extract_docx_images()
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp']:
            return '', [self.file.read()]
        elif ext in ['.csv', '.xlsx']:
            return self._extract_table_text(), []
        else:
            return self.file.read().decode('utf-8'), []

    def _extract_pdf_text(self):
        self.file.seek(0)
        text = []
        with pdfplumber.open(self.file) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return '\n'.join(text)

    def _extract_pdf_images(self):
        self.file.seek(0)
        images = []
        with pdfplumber.open(self.file) as pdf:
            for page in pdf.pages:
                for img in page.images:
                    try:
                        im = page.to_image(resolution=150)
                        cropped = im.original.crop((img['x0'], img['top'], img['x1'], img['bottom']))
                        buf = io.BytesIO()
                        cropped.save(buf, format='PNG')
                        images.append(buf.getvalue())
                    except Exception:
                        continue
        return images

    def _extract_docx_text(self):
        self.file.seek(0)
        doc = docx.Document(self.file)
        lines = []
        for para in doc.paragraphs:
            lines.append(para.text)
        return '\n'.join(lines)

    def _extract_docx_images(self):
        self.file.seek(0)
        doc = docx.Document(self.file)
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    img_bytes = rel.target_part.blob
                    images.append(img_bytes)
                except Exception:
                    continue
        return images

    def _extract_table_text(self):
        ext = os.path.splitext(self.file.name)[1].lower()
        self.file.seek(0)
        try:
            if ext == '.csv':
                df = pd.read_csv(self.file)
            elif ext == '.xlsx':
                df = pd.read_excel(self.file)
            else:
                return ''
        except Exception:
            return ''
        # Her satırı "tablo başlığı: veri" formatında birleştir
        lines = []
        header = df.columns.tolist()
        for idx, row in df.iterrows():
            row_str = ', '.join([f"{col}: {row[col]}" for col in header])
            lines.append(row_str)
        return '\n'.join(lines)

    def chunk_text(self, text, max_tokens=400):
        # Paragraflara böl, başlıkları tespit et, chunk'lara başlık ekle
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        tokenizer = tiktoken.get_encoding('cl100k_base')
        chunks = []
        current_chunk = []
        current_tokens = 0
        current_heading = None
        heading_map = {}
        # Başlık tespiti
        for i, para in enumerate(paragraphs):
            heading = self._detect_heading(para)
            if heading:
                current_heading = heading
            heading_map[i] = current_heading
        # Chunk'lama
        for i, para in enumerate(paragraphs):
            tokens = len(tokenizer.encode(para))
            if current_tokens + tokens > max_tokens and current_chunk:
                chunks.append(current_chunk)
                current_chunk = []
                current_tokens = 0
            current_chunk.append((i, para))
            current_tokens += tokens
        if current_chunk:
            chunks.append(current_chunk)
        # Her chunk'a başlık ekle
        chunk_headings = []
        for chunk in chunks:
            # Chunk'taki ilk paragrafın başlığı
            first_para_idx = chunk[0][0]
            heading = heading_map.get(first_para_idx)
            chunk_headings.append(heading)
        return list(zip(chunks, chunk_headings))

    def _detect_heading(self, para):
        # DOCX için heading stilleriyle tespit zaten yapılır, burada düz metin/PDF için heuristik
        # Markdown başlıkları
        if para.startswith('#'):
            return para.lstrip('#').strip()
        # Büyük harfli, kısa satır (ör: BÖLÜM 1, GİRİŞ)
        if para.isupper() and 3 < len(para) < 60:
            return para
        # Numara ile başlayan başlıklar (ör: 1. Giriş, 2.1 Alt Başlık)
        if (para[:3].replace('.', '').isdigit() and len(para) < 80):
            return para
        # Son olarak: 30 karakterden kısa, ilk harfi büyük, sonunda nokta yoksa
        if len(para) < 50 and para[:1].isupper() and not para.endswith('.'):
            return para
        return None

    def chunk_table(self, text, max_rows=10):
        # Tabloyu satır bazında chunk'la
        lines = [l for l in text.split('\n') if l.strip()]
        chunks = []
        for i in range(0, len(lines), max_rows):
            chunk = lines[i:i+max_rows]
            chunks.append((chunk, i, min(i+max_rows-1, len(lines)-1)))
        return chunks

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
        QDRANT_FILE_COLLECTION = 'gpt_package_files'
        try:
            # Koleksiyon kontrolü semantic.py'ye taşındı
            ensure_collection_exists(QDRANT_FILE_COLLECTION, vector_size=384)
        except Exception:
            pass
        try:
            text, images = self.get_file_content()
        except Exception:
            return
        file_name = self.file.name
        title = self.description or file_name
        points = []
        ext = os.path.splitext(self.file.name)[1].lower()
        # Tablo dosyası ise özel chunk ve embedding
        if ext in ['.csv', '.xlsx']:
            table_chunks = self.chunk_table(text)
            for chunk_idx, (chunk_lines, start_row, end_row) in enumerate(table_chunks):
                chunk_text = '\n'.join(chunk_lines)
                embedding = get_embedding(chunk_text)
                points.append(
                    qdrant_models.PointStruct(
                        id=f"{self.id}_table_{chunk_idx}",
                        vector=embedding,
                        payload={
                            "gpt_package_file_id": str(self.id),
                            "gpt_package_id": str(self.gpt_package.id),
                            "file_name": file_name,
                            "title": title,
                            "chunk_index": chunk_idx,
                            "row_start": start_row,
                            "row_end": end_row,
                            "type": "table",
                        }
                    )
                )
        else:
            # Metin chunk'ları
            chunk_heading_pairs = self.chunk_text(text)
            for chunk_idx, (chunk, heading) in enumerate(chunk_heading_pairs):
                chunk_text = '\n'.join([p[1] for p in chunk])
                embedding = get_embedding(chunk_text)
                para_indices = [p[0] for p in chunk]
                points.append(
                    qdrant_models.PointStruct(
                        id=f"{self.id}_text_{chunk_idx}",
                        vector=embedding,
                        payload={
                            "gpt_package_file_id": str(self.id),
                            "gpt_package_id": str(self.gpt_package.id),
                            "file_name": file_name,
                            "title": title,
                            "chunk_index": chunk_idx,
                            "paragraph_indices": para_indices,
                            "heading": heading,
                            "type": "text",
                        }
                    )
                )
            # Görsel chunk'ları
            if images:
                clip_model, clip_preprocess = get_clip_model()
                import torch
                for img_idx, img_bytes in enumerate(images):
                    try:
                        image = Image.open(io.BytesIO(img_bytes)).convert('RGB')
                        image_input = clip_preprocess(image).unsqueeze(0)
                        with torch.no_grad():
                            image_features = clip_model.encode_image(image_input).squeeze(0).cpu().numpy().tolist()
                        points.append(
                            qdrant_models.PointStruct(
                                id=f"{self.id}_image_{img_idx}",
                                vector=image_features,
                                payload={
                                    "gpt_package_file_id": str(self.id),
                                    "gpt_package_id": str(self.gpt_package.id),
                                    "file_name": file_name,
                                    "title": title,
                                    "chunk_index": img_idx,
                                    "type": "image",
                                }
                            )
                        )
                    except Exception:
                        continue
        if points:
            for point in points:
                add_to_qdrant(QDRANT_FILE_COLLECTION, "", point.payload | {"id": point.id})

    def delete(self, *args, **kwargs):
        QDRANT_FILE_COLLECTION = 'gpt_package_files'
        try:
            ext = os.path.splitext(self.file.name)[1].lower()
            if ext in ['.csv', '.xlsx']:
                text, _ = self.get_file_content()
                table_chunks = self.chunk_table(text)
                point_ids = [f"{self.id}_table_{i}" for i in range(len(table_chunks))]
            else:
                text, images = self.get_file_content()
                chunk_heading_pairs = self.chunk_text(text)
                point_ids = [f"{self.id}_text_{i}" for i in range(len(chunk_heading_pairs))]
                point_ids += [f"{self.id}_image_{i}" for i in range(len(images))]
            delete_from_qdrant(QDRANT_FILE_COLLECTION, point_ids)
        except Exception:
            pass
        super().delete(*args, **kwargs)

